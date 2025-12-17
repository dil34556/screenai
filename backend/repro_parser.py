from docx import Document
import os
import sys

# Ensure backend/screenai is in path if needed, usually running from backend/ works
sys.path.append(os.getcwd())

try:
    doc = Document()
    doc.add_paragraph("John Doe")
    doc.add_paragraph("Email: john@example.com")
    doc.add_paragraph("Phone: 555-0199")
    doc.add_paragraph("Experience: 5 years in Python and Django.")
    doc.save("test_resume.docx")
    print("Created test_resume.docx")
except Exception as e:
    print(f"Failed to create docx: {e}")
    sys.exit(1)

try:
    # FORCE INVALID KEY to test Regex Fallback
    os.environ["GOOGLE_API_KEY"] = "INVALID_KEY_FOR_TESTING_FALLBACK"
    os.environ["GEMINI_API_KEY"] = "INVALID_KEY_FOR_TESTING_FALLBACK"
    
    from screenai.services.resume_parser.parser import parse_resume
    print("Imported parse_resume")
    
    result = parse_resume(os.path.abspath("test_resume.docx"))
    print("Parsed successfully (via Fallback)!")
    print(result)
    
    # VERIFY STRUCTURE
    if "data" in result and "email" in result["data"]:
        print("SUCCESS: Result has 'data' key and correct structure.")
    else:
        print("FAILURE: Result missing 'data' key!")

except Exception as e:
    print(f"PARSING FAILED: {e}")
    import traceback
    traceback.print_exc()
