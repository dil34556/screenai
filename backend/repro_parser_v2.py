from docx import Document
import os
import sys
import datetime

# Ensure backend/screenai is in path if needed, usually running from backend/ works
sys.path.append(os.getcwd())

try:
    doc = Document()
    doc.add_paragraph("Alice Developer")
    doc.add_paragraph("Email: alice@example.com")
    doc.add_paragraph("Phone: 555-0199")
    doc.add_paragraph("SKILLS: Python, React, Docker, Kubernetes")
    doc.add_paragraph("EXPERIENCE")
    doc.add_paragraph("Senior Engineer at Google | 2020 - Present")
    doc.add_paragraph("Developer at Startup Inc | 2018 - 2020")
    doc.save("test_resume_v2.docx")
    print("Created test_resume_v2.docx")
except Exception as e:
    print(f"Failed to create docx: {e}")
    sys.exit(1)

try:
    # FORCE INVALID KEY to test Regex Fallback
    os.environ["GOOGLE_API_KEY"] = "INVALID_KEY_FOR_TESTING_FALLBACK"
    os.environ["GEMINI_API_KEY"] = "INVALID_KEY_FOR_TESTING_FALLBACK"
    
    from screenai.services.resume_parser.parser import parse_resume
    print("Imported parse_resume")
    
    result = parse_resume(os.path.abspath("test_resume_v2.docx"))
    print("Parsed result:")
    print(result)
    
    data = result.get('data', {})
    skills = data.get('skills', [])
    exp = data.get('work_experience', [])
    
    print("\n--- VERIFICATION ---")
    if "Python" in skills and "React" in skills:
        print("SUCCESS: Skills extracted!")
    else:
        print(f"FAILURE: Skills missing. Found: {skills}")
        
    if len(exp) >= 2:
        print(f"SUCCESS: Found {len(exp)} experience items.")
        print(exp)
    else:
         print(f"FAILURE: Experience missing. Found: {exp}")

except Exception as e:
    print(f"PARSING FAILED: {e}")
    import traceback
    traceback.print_exc()
