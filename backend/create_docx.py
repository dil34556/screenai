from docx import Document

document = Document()
document.add_heading('John Doe', 0)
document.add_paragraph('Email: john.doe@example.com | Phone: 555-0100')
document.add_heading('Experience', level=1)
document.add_paragraph('Software Engineer at Tech Corp (2020-Present)')
document.add_heading('Education', level=1)
document.add_paragraph('BS Computer Science, University of Examples')

document.save('test_resume.docx')
print("Created test_resume.docx")
