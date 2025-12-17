from docx import Document
import os
import sys

# Ensure backend/screenai is in path if needed
sys.path.append(os.getcwd())

try:
    doc = Document()
    doc.add_paragraph("Bob Builder")
    doc.add_paragraph("Email: bob@example.com")
    doc.add_paragraph("Phone: 555-0199")
    
    # 1. Section Header
    doc.add_paragraph("SKILLS")
    # 2. Content (Comma sep) with new keywords
    doc.add_paragraph("Flutter, Dart, PowerBI, Machine Learning, Microsoft Excel")
    # 3. Content (Bullets? Just lines here)
    doc.add_paragraph("Solving Complex Problems")
    
    doc.add_paragraph("EDUCATION")
    doc.add_paragraph("University of Data")
    
    doc.save("test_resume_v3.docx")
    print("Created test_resume_v3.docx")
except Exception as e:
    print(f"Failed to create docx: {e}")
    sys.exit(1)

try:
    # FORCE INVALID KEY to test Regex Fallback
    os.environ["GOOGLE_API_KEY"] = "INVALID_KEY_FOR_TESTING_FALLBACK"
    os.environ["GEMINI_API_KEY"] = "INVALID_KEY_FOR_TESTING_FALLBACK"
    
    from screenai.services.resume_parser.parser import parse_resume
    
    result = parse_resume(os.path.abspath("test_resume_v3.docx"))
    print("Parsed result:")
    skills = result.get('data', {}).get('skills', [])
    print(f"Skills Found: {skills}")
    
    # Verify strict targets
    targets = ["Flutter", "PowerBI", "Machine Learning"]
    found_all = True
    for t in targets:
        # Check if t is in skills list (case insensitive or exact)
        # Our parser titles them, so "Flutter" should match "Flutter"
        if not any(s.lower() == t.lower() for s in skills):
            print(f"MISSING: {t}")
            found_all = False
            
    if found_all:
        print("SUCCESS: All target skills found via fallback!")
    else:
        print("FAILURE: Some skills missing.")

except Exception as e:
    print(f"PARSING FAILED: {e}")
    import traceback
    traceback.print_exc()
